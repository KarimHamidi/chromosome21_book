# -*- coding: utf-8 -*-
"""
Created on Thu Oct 20 18:57:41 2016

@author: karim
"""
#Définition du répertoire qui sera utilisé
import os 
os.chdir('C:\\Users\\karim\\FirstStepProject\\.spyproject')
#importation de Gzip et de Requests permettant le téléchargement des séquences
import requests
import gzip
import requests
#entré tous les URLs des séquences voulus dans la liste

list_url = ['http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_GL383578v2_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_GL383579v2_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_GL383580v2_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_GL383581v2_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_KI270872v1_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_KI270873v1_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21_KI270874v1_alt.fa.gz','http://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/chr21.fa.gz']

for i in list_url:
    filename = i.split("/")[-1]
    with open(filename, "wb") as f:
        r = requests.get(i)
        f.write(r.content)
#décompresser les fichiers et les sauver dans le meme directory   
import os      
import gzip
import glob
import os.path
source_dir = 'C:\\Users\\karim\\FirstStepProject\\.spyproject'
dest_dir = 'C:\\Users\\karim\\FirstStepProject\\.spyproject'

for src_name in glob.glob(os.path.join(source_dir, '*.gz')):
    base = os.path.basename(src_name)
    dest_name = os.path.join(dest_dir, base[:-3])
    with gzip.open(src_name, 'rb') as infile:
        with open(dest_name, 'wb') as outfile:
            for line in infile:
                outfile.write(line)


        

#passer par une string pour créer un fichier texte. Le fichier sequence.txt contient la séquence   
seqence_essai 
import textwrap 
contig1_string = str(contig1_content,'utf-8') #type(file_string) Out[19]: str permet de passer d'un 8bytà un string simple (variante du meme objet mais plus manipulable)
contig1_string2 = file_string[file_string.find('\n')+1:] #la première ligne content l'id est supprimée
contig1_string3 = file_string2.lower()
contig1_string4 = file_string3.replace('\n','')
contig1_string5 = textwrap.fill(file_string4,80)

#insérer la séquence dans un fichier texte                           
text_file = open("contig1.txt", "w")
text_file.write(file_string5)
text_file.close()



















#GENERATION DU PDF :
from reportlab import *
#Génération des pages suivantes platypus? Text?
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm

#génération d'un pdf contenant la séquence
def generate_book(c):
    title = 'Chromosme 21'
    ptr = open("Seqmin80.txt", "r")  #fichier text à convertir
    lines = ptr.readlines()
    ptr.close()
    i = 800
    numeroLine = 0
    c = canvas.Canvas('Chr21C60999.pdf')
    image = 'Human_chromosome_21_from_Gene_Gateway_-_no_label.png'
    c.setFont('Helvetica',60, leading= None)
    c.drawCentredString(300, 800,title)
    c.drawImage(image,100,100,width=200, height =200)
    c.showPage()
    while numeroLine < len(lines):
        if numeroLine - len(lines) < 63: # Défini les nombre de ligne par page
            i=800
            for line in lines[numeroLine:numeroLine+63]:
                c.setFont('Helvetica',10,leading=None)
                c.drawString(15, i, line.strip())   #si centré 300 sinon 15
                numeroLine += 1
                i -= 12                 #espace interligne
            c.showPage()
#    else:
#        i = 800
#        for line in lines[numeroLine:]:
#           numeroLine += 1
#           i -= 12

c.save()

generate_book(c)

#doc X pour créer un fichier word ce qui pourrait etre une solution alternative à REPORTLAB
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)
seq = 'bsjdjedklwdjdwedwdcjweiocjweijwicjeijcwioecjiweocdjewio\n djadkljad'
p = document.add_paragraph(seq)
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('Human_chromosome_21_from_Gene_Gateway_-_no_label.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

document.add_page_break()

document.save('demo.docx')



    